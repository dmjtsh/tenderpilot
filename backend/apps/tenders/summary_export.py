import os
from io import BytesIO
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

FONT_DIR = os.path.join(os.path.dirname(__file__), "fonts")

# Item value is either a string (key-value pair) or a list of strings (subheading + bullets).
SectionItem = tuple[str, str | list[str]]
Section = tuple[str, list[SectionItem]]


def _fmt_money(val: float | int | None) -> str:
    if val is None:
        return ""
    if val >= 1_000_000:
        return f"{val / 1_000_000:,.2f} млн ₽".replace(",", " ")
    if val >= 1_000:
        return f"{val / 1_000:,.1f} тыс. ₽".replace(",", " ")
    return f"{val:,.0f} ₽".replace(",", " ")


def _fmt_pct(val: float | int | None) -> str:
    if val is None:
        return ""
    return f"{val}%"


def _severity_label(s: str) -> str:
    return {"high": "высокий", "medium": "средний", "low": "низкий"}.get(s, s)


def _collect(items: list[str]) -> list[str] | None:
    """Return list only if non-empty, else None (for easy skip)."""
    return items if items else None


def _sections(summary: dict) -> list[Section]:
    sections: list[Section] = []

    # ── Общая информация ────────────────────────────────────────────────
    gen = summary.get("general") or {}
    items: list[SectionItem] = []
    if gen.get("title"):
        items.append(("Название", gen["title"]))
    if gen.get("customer_name"):
        items.append(("Заказчик", gen["customer_name"]))
    if gen.get("customer_inn"):
        items.append(("ИНН", gen["customer_inn"]))
    if gen.get("law_type"):
        items.append(("Закон", gen["law_type"]))
    if gen.get("nmck") is not None:
        items.append(("НМЦК", _fmt_money(gen["nmck"])))
    if gen.get("region"):
        items.append(("Регион", gen["region"]))
    if items:
        sections.append(("Общая информация", items))

    # ── Анализ заказчика ────────────────────────────────────────────────
    ca = summary.get("customer_analysis")
    if ca:
        items = []
        if ca.get("name"):
            items.append(("Организация", ca["name"]))
        if ca.get("inn"):
            items.append(("ИНН", ca["inn"]))
        if ca.get("org_type"):
            items.append(("Тип", ca["org_type"]))
        if ca.get("industry"):
            items.append(("Отрасль", ca["industry"]))
        if ca.get("region"):
            items.append(("Регион", ca["region"]))
        if ca.get("founded_date"):
            items.append(("Дата основания", ca["founded_date"]))
        if ca.get("reliability"):
            items.append(("Надёжность", ca["reliability"]))
        fin = ca.get("financials") or {}
        if fin.get("revenue_rub"):
            year = f" ({fin['revenue_year']})" if fin.get("revenue_year") else ""
            items.append(("Выручка", f"{_fmt_money(fin['revenue_rub'])}{year}"))
        if fin.get("profit_rub"):
            items.append(("Прибыль", _fmt_money(fin["profit_rub"])))
        if fin.get("employees_count"):
            items.append(("Сотрудников", str(fin["employees_count"])))
        ph = ca.get("procurement_history") or {}
        if ph.get("total_purchases"):
            items.append(("Закупок", str(ph["total_purchases"])))
        if ph.get("total_amount_rub"):
            items.append(("Общая сумма закупок", _fmt_money(ph["total_amount_rub"])))
        if ph.get("avg_contract_rub"):
            items.append(("Средний контракт", _fmt_money(ph["avg_contract_rub"])))
        ri = ca.get("risk_indicators") or {}
        flags = ri.get("red_flags") or []
        if flags:
            items.append(("Красные флаги", flags))
        if ri.get("licenses_count"):
            items.append(("Лицензий", str(ri["licenses_count"])))
        if ri.get("arbitration_count"):
            items.append(("Арбитражных дел", str(ri["arbitration_count"])))
        if ri.get("fssp_count"):
            items.append(("Исполнительных производств", str(ri["fssp_count"])))
        notes = ca.get("notes") or []
        if notes:
            items.append(("Примечания", notes))
        if items:
            sections.append(("Анализ заказчика", items))

    # ── Описание работ ──────────────────────────────────────────────────
    wd = summary.get("work_description")
    if wd:
        items = []
        if wd.get("subject"):
            items.append(("Предмет", wd["subject"]))
        if wd.get("tender_type"):
            items.append(("Тип тендера", wd["tender_type"]))
        scope = wd.get("scope") or {}
        acts = _collect(scope.get("main_activities") or [])
        if acts:
            items.append(("Состав работ", acts))
        delivs = _collect(scope.get("deliverables") or [])
        if delivs:
            items.append(("Результаты поставки", delivs))
        for vm in (scope.get("volume_metrics") or []):
            items.append((vm.get("metric", "Объём"), str(vm.get("value", ""))))
        loc = wd.get("location") or {}
        if loc.get("address"):
            items.append(("Адрес", loc["address"]))
        elif loc.get("region"):
            items.append(("Регион работ", loc["region"]))
        ts = wd.get("technical_specs") or {}
        equip = _collect(ts.get("equipment") or [])
        if equip:
            items.append(("Оборудование", equip))
        soft = _collect(ts.get("software") or [])
        if soft:
            items.append(("Программное обеспечение", soft))
        stds = _collect(ts.get("standards") or [])
        if stds:
            items.append(("Стандарты", stds))
        mats = _collect(ts.get("materials") or [])
        if mats:
            items.append(("Материалы", mats))
        acc = wd.get("acceptance") or {}
        if acc.get("procedure"):
            items.append(("Приёмка", acc["procedure"]))
        if acc.get("review_days"):
            items.append(("Срок приёмки", f"{acc['review_days']} дн."))
        if acc.get("warranty_months"):
            items.append(("Гарантия", f"{acc['warranty_months']} мес."))
        if acc.get("support_after_delivery"):
            items.append(("Поддержка после поставки", acc["support_after_delivery"]))
        sub = wd.get("subcontracting") or {}
        if sub.get("allowed") is not None:
            val = "Да" if sub["allowed"] else "Нет"
            if sub.get("max_pct"):
                val += f" (до {sub['max_pct']}%)"
            items.append(("Субподряд", val))
        if sub.get("restrictions"):
            items.append(("Ограничения субподряда", sub["restrictions"]))
        if items:
            sections.append(("Описание работ", items))

    # ── Финансы ─────────────────────────────────────────────────────────
    fin = summary.get("financial")
    if fin:
        items = []
        adv = fin.get("advance") or {}
        if adv.get("has_advance"):
            desc = "Да"
            if adv.get("amount_pct"):
                desc += f" ({adv['amount_pct']}%)"
            if adv.get("amount_rub"):
                desc += f" — {_fmt_money(adv['amount_rub'])}"
            items.append(("Аванс", desc))
            if adv.get("description"):
                items.append(("Условия аванса", adv["description"]))
        elif adv.get("has_advance") is False:
            items.append(("Аванс", "Не предусмотрен"))
        pay = fin.get("payment") or {}
        if pay.get("structure"):
            items.append(("Оплата", pay["structure"]))
        if pay.get("schedule_description"):
            items.append(("График оплаты", pay["schedule_description"]))
        if pay.get("term_days"):
            dtype = pay.get("term_days_type") or "дн."
            items.append(("Срок оплаты", f"{pay['term_days']} {dtype}"))
        if pay.get("count"):
            items.append(("Количество платежей", str(pay["count"])))
        if pay.get("fixed_price") is not None:
            items.append(("Фиксированная цена", "Да" if pay["fixed_price"] else "Нет"))
        if fin.get("funding_source"):
            items.append(("Источник", fin["funding_source"]))
        sec = fin.get("securities") or {}
        if sec.get("bid_amount_rub") or sec.get("bid_pct"):
            val = _fmt_money(sec.get("bid_amount_rub")) or _fmt_pct(sec.get("bid_pct"))
            items.append(("Обеспечение заявки", val))
        if sec.get("contract_amount_rub") or sec.get("contract_pct"):
            val = _fmt_money(sec.get("contract_amount_rub")) or _fmt_pct(sec.get("contract_pct"))
            if sec.get("contract_form"):
                val += f" ({sec['contract_form']})"
            items.append(("Обеспечение контракта", val))
        if sec.get("return_term_days"):
            items.append(("Возврат обеспечения", f"{sec['return_term_days']} дн."))
        pen = fin.get("penalties") or {}
        if pen.get("delay_pen_formula"):
            base = f" (от {pen['delay_pen_base']})" if pen.get("delay_pen_base") else ""
            items.append(("Пеня за просрочку", f"{pen['delay_pen_formula']}{base}"))
        if pen.get("fixed_fine_rub"):
            basis = f" ({pen['fixed_fine_basis']})" if pen.get("fixed_fine_basis") else ""
            items.append(("Штраф", f"{_fmt_money(pen['fixed_fine_rub'])}{basis}"))
        ad = fin.get("antidumping") or {}
        if ad.get("applicable"):
            desc = f"Порог {ad.get('threshold_pct', '?')}%"
            if ad.get("multiplier"):
                desc += f", множитель ×{ad['multiplier']}"
            items.append(("Антидемпинг", desc))
        if fin.get("cash_flow_note"):
            items.append(("Примечание", fin["cash_flow_note"]))
        if items:
            sections.append(("Финансы", items))

    # ── Сроки ───────────────────────────────────────────────────────────
    tl = summary.get("timeline")
    if tl:
        items = []
        td = tl.get("total_duration") or {}
        if td.get("days"):
            dtype = td.get("day_type") or "дн."
            items.append(("Срок", f"{td['days']} {dtype}"))
        if td.get("description"):
            items.append(("", td["description"]))
        stages = tl.get("stages") or []
        if stages:
            stage_lines = []
            for st in stages:
                name = st.get("name", "")
                parts = []
                if st.get("duration_days"):
                    parts.append(f"{st['duration_days']} дн.")
                if st.get("start_date"):
                    parts.append(f"с {st['start_date']}")
                if st.get("end_date"):
                    parts.append(f"по {st['end_date']}")
                suffix = f" ({', '.join(parts)})" if parts else ""
                line = f"Этап {st.get('number', '')}: {name}{suffix}"
                stage_lines.append(line)
                for dlv in (st.get("deliverables") or []):
                    stage_lines.append(f"  → {dlv}")
            items.append(("Этапы", stage_lines))
        kd = tl.get("key_dates") or {}
        if kd.get("submission_deadline"):
            items.append(("Подача заявок", kd["submission_deadline"]))
        if kd.get("auction_date"):
            items.append(("Аукцион", kd["auction_date"]))
        if kd.get("contract_sign_deadline_days"):
            items.append(("Срок подписания контракта", f"{kd['contract_sign_deadline_days']} дн."))
        if kd.get("warranty_months"):
            items.append(("Гарантия", f"{kd['warranty_months']} мес."))
        if tl.get("urgency_note"):
            items.append(("Срочность", tl["urgency_note"]))
        if items:
            sections.append(("Сроки", items))

    # ── Требования к участникам ─────────────────────────────────────────
    req = summary.get("requirements")
    if req:
        items = []
        elig = req.get("eligibility") or {}

        # Допуск
        access_lines = []
        for lic in (elig.get("licenses") or []):
            name = lic.get("name", "")
            extra = []
            if lic.get("issuer"):
                extra.append(lic["issuer"])
            if lic.get("mandatory") is False:
                extra.append("необяз.")
            suffix = f" ({', '.join(extra)})" if extra else ""
            access_lines.append(f"Лицензия: {name}{suffix}")
        sro = elig.get("sro") or {}
        if sro.get("required"):
            access_lines.append(f"СРО: {sro.get('type') or 'Требуется'}")
        exp = elig.get("experience") or {}
        if exp.get("description"):
            access_lines.append(f"Опыт: {exp['description']}")
        elif exp.get("min_contracts"):
            access_lines.append(f"Опыт: мин. {exp['min_contracts']} контрактов")
        if exp.get("years"):
            access_lines.append(f"Опыт: {exp['years']} лет")
        if exp.get("min_amount_rub"):
            access_lines.append(f"Мин. сумма опыта: {_fmt_money(exp['min_amount_rub'])}")
        for st in (elig.get("staff") or []):
            quals = f", {st['qualifications']}" if st.get("qualifications") else ""
            access_lines.append(f"Персонал: {st.get('role', '')} — {st.get('count', '?')} чел.{quals}")
        efin = elig.get("financial") or {}
        if efin.get("min_revenue_rub"):
            access_lines.append(f"Мин. выручка: {_fmt_money(efin['min_revenue_rub'])}")
        if efin.get("no_bankruptcy"):
            access_lines.append("Отсутствие банкротства")
        if efin.get("no_tax_debt"):
            access_lines.append("Отсутствие налоговой задолженности")
        for o in (elig.get("other") or []):
            access_lines.append(o)
        if access_lines:
            items.append(("Допуск участника", access_lines))

        # Подача заявки
        sub = req.get("submission") or {}
        doc_lines = []
        for form in (sub.get("forms") or []):
            doc_lines.append(f"Форма: {form}")
        if sub.get("electronic_signature"):
            doc_lines.append("Электронная подпись (ЭЦП) — требуется")
        for doc in (sub.get("documents") or []):
            m = "обяз." if doc.get("mandatory") else "необяз."
            doc_lines.append(f"{doc.get('name', '')} ({m})")
        if doc_lines:
            items.append(("Документы для подачи", doc_lines))

        # Критерии оценки
        for ec in (req.get("evaluation_criteria") or []):
            w = f" ({ec['weight_pct']}%)" if ec.get("weight_pct") else ""
            desc = f" — {ec['description']}" if ec.get("description") else ""
            items.append(("Критерий оценки", f"{ec.get('name', '')}{w}{desc}"))

        # Ограничения
        restr_lines = []
        restr = req.get("restrictions") or {}
        if restr.get("smp_only"):
            restr_lines.append("Только для СМП")
        if restr.get("national_regime"):
            restr_lines.append("Национальный режим")
        if restr.get("region_restriction"):
            restr_lines.append(restr["region_restriction"])
        for o in (restr.get("other") or []):
            restr_lines.append(o)
        if restr_lines:
            items.append(("Ограничения", restr_lines))

        if req.get("notes"):
            items.append(("Примечание", req["notes"]))
        if items:
            sections.append(("Требования к участникам", items))

    # ── Риски ───────────────────────────────────────────────────────────
    risks = summary.get("risks")
    if risks:
        items = []
        if risks.get("overall_risk"):
            items.append(("Общий уровень", risks["overall_risk"]))
        if risks.get("risk_summary"):
            items.append(("Резюме", risks["risk_summary"]))
        for cat, label in [
            ("certification_risks", "Сертификация"),
            ("financial_risks", "Финансы"),
            ("technical_risks", "Техника"),
            ("legal_risks", "Правовые"),
            ("timeline_risks", "Сроки"),
        ]:
            cat_risks = risks.get(cat) or []
            if cat_risks:
                risk_lines = []
                for r in cat_risks:
                    sev = _severity_label(r.get("severity", ""))
                    risk_lines.append(f"[{sev}] {r.get('risk', '')}")
                items.append((label, risk_lines))
        unusual = risks.get("unusual_conditions") or []
        if unusual:
            items.append(("Нетипичные условия", unusual))
        if items:
            sections.append(("Риски", items))

    return sections


# ── DOCX ─────────────────────────────────────────────────────────────────────

def render_summary_docx(tender: Any, summary: dict) -> BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)

    h = doc.add_heading(level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.add_run(f"AI-резюме тендера №{tender.number}")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_p.add_run(tender.title or "")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_paragraph("")

    for section_title, items in _sections(summary):
        h = doc.add_heading(section_title, level=2)
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x11, 0x18, 0x27)

        for label, value in items:
            if isinstance(value, list):
                if label:
                    p = doc.add_paragraph()
                    run = p.add_run(label)
                    run.bold = True
                    run.font.size = Pt(10)
                for line in value:
                    p = doc.add_paragraph(style="List Bullet")
                    run = p.add_run(str(line))
                    run.font.size = Pt(10)
            else:
                p = doc.add_paragraph()
                if label:
                    run = p.add_run(f"{label}: ")
                    run.bold = True
                    run.font.size = Pt(10)
                run = p.add_run(str(value))
                run.font.size = Pt(10)

        doc.add_paragraph("")

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── PDF ──────────────────────────────────────────────────────────────────────

def _register_fonts():
    if "DejaVuSans" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(TTFont("DejaVuSans", os.path.join(FONT_DIR, "DejaVuSans.ttf")))
        pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", os.path.join(FONT_DIR, "DejaVuSans-Bold.ttf")))


def _esc(text: str) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _get_styles():
    _register_fonts()
    styles = getSampleStyleSheet()

    styles.add(ParagraphStyle(
        "TitleRu", parent=styles["Title"],
        fontName="DejaVuSans-Bold", fontSize=16, leading=20,
        textColor=HexColor("#111827"), alignment=0,
        spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "SubtitleRu", parent=styles["Normal"],
        fontName="DejaVuSans", fontSize=11, leading=14,
        textColor=HexColor("#4B5563"), spaceAfter=12,
    ))
    styles.add(ParagraphStyle(
        "SectionRu", parent=styles["Heading2"],
        fontName="DejaVuSans-Bold", fontSize=13, leading=16,
        textColor=HexColor("#111827"), spaceBefore=14, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "BodyRu", parent=styles["Normal"],
        fontName="DejaVuSans", fontSize=10, leading=13,
        textColor=HexColor("#1F2937"), spaceAfter=3,
    ))
    styles.add(ParagraphStyle(
        "SubHeadRu", parent=styles["Normal"],
        fontName="DejaVuSans-Bold", fontSize=10, leading=13,
        textColor=HexColor("#111827"), spaceBefore=6, spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "BulletRu", parent=styles["Normal"],
        fontName="DejaVuSans", fontSize=10, leading=13,
        textColor=HexColor("#1F2937"), spaceAfter=2,
        leftIndent=14, bulletIndent=0,
    ))
    return styles


def render_summary_pdf(tender: Any, summary: dict) -> BytesIO:
    buf = BytesIO()
    styles = _get_styles()

    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=2 * cm, bottomMargin=2 * cm,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
    )

    story: list = []

    story.append(Paragraph(f"AI-резюме тендера №{tender.number}", styles["TitleRu"]))
    story.append(Paragraph(tender.title or "", styles["SubtitleRu"]))
    story.append(Spacer(1, 6))

    for section_title, items in _sections(summary):
        story.append(HRFlowable(width="100%", thickness=0.5, color=HexColor("#E5E7EB"), spaceAfter=4))
        story.append(Paragraph(section_title, styles["SectionRu"]))

        for label, value in items:
            if isinstance(value, list):
                if label:
                    story.append(Paragraph(_esc(label), styles["SubHeadRu"]))
                for line in value:
                    story.append(Paragraph(f"\u2022 {_esc(line)}", styles["BulletRu"]))
            else:
                value_esc = _esc(value)
                if label:
                    label_esc = _esc(label)
                    text = f'<font name="DejaVuSans-Bold">{label_esc}:</font> {value_esc}'
                else:
                    text = value_esc
                story.append(Paragraph(text, styles["BodyRu"]))

        story.append(Spacer(1, 6))

    doc.build(story)
    buf.seek(0)
    return buf
